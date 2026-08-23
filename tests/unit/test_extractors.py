"""Unit tests for multi-format text and structure extractors."""

import io
from pathlib import Path

import docx
import pytest

from app.core.exceptions import ValidationError
from app.rag.ingestion.extractors import (
    DocxExtractor,
    MarkdownExtractor,
    PDFExtractor,
    PlainTextExtractor,
    get_extractor_for_type,
)


def test_plain_text_extractor():
    """Verify plain text extraction and paragraph segmentation."""
    extractor = PlainTextExtractor()
    content = b"Paragraph 1 content.\n\nParagraph 2 with details."
    doc = extractor.extract(content, "my_notes.txt")

    assert doc.title == "My Notes"
    assert "Paragraph 1 content" in doc.full_text
    assert len(doc.sections) == 2
    assert doc.sections[0].heading == "Paragraph 1"
    assert doc.sections[1].heading == "Paragraph 2"


def test_markdown_extractor_headers():
    """Verify markdown extraction preserves headings and hierarchical sections."""
    extractor = MarkdownExtractor()
    md_content = b"""# Main Architecture
System overview description.

## Section 1: Ingestion
Ingestion details and data loading.

## Section 2: Retrieval
Vector search and BM25 details.
"""
    doc = extractor.extract(md_content, "architecture_spec.md")

    assert doc.title == "Main Architecture"
    assert len(doc.sections) >= 2
    headings = [s.heading for s in doc.sections]
    assert "Section 1: Ingestion" in headings
    assert "Section 2: Retrieval" in headings


def test_pdf_extractor_with_real_sample():
    """Verify PDF extraction with sample PDF preserving pages."""
    pdf_path = Path("data/raw/sample_paper.pdf")
    assert pdf_path.exists()

    with open(pdf_path, "rb") as f:
        content = f.read()

    extractor = PDFExtractor()
    doc = extractor.extract(content, "sample_paper.pdf")

    assert doc.total_pages == 2
    assert len(doc.sections) == 2
    assert doc.sections[0].page_number == 1
    assert doc.sections[1].page_number == 2
    assert "SentinelRAG" in doc.sections[0].content
    assert "Hybrid Retrieval" in doc.sections[1].content


def test_docx_extractor():
    """Verify DOCX extractor parses headings and paragraphs."""
    doc_io = io.BytesIO()
    doc_obj = docx.Document()
    doc_obj.add_heading("Project Overview", level=1)
    doc_obj.add_paragraph("SentinelRAG project overview paragraph.")
    doc_obj.add_heading("Methodology", level=2)
    doc_obj.add_paragraph("Methodology paragraph details.")
    doc_obj.save(doc_io)
    docx_bytes = doc_io.getvalue()

    extractor = DocxExtractor()
    doc = extractor.extract(docx_bytes, "project_doc.docx")

    assert "Project Overview" in doc.full_text
    assert "Methodology" in doc.full_text
    assert len(doc.sections) >= 2


def test_get_extractor_for_type():
    """Verify extractor factory resolution."""
    assert isinstance(get_extractor_for_type("pdf"), PDFExtractor)
    assert isinstance(get_extractor_for_type("markdown"), MarkdownExtractor)
    assert isinstance(get_extractor_for_type("text"), PlainTextExtractor)
    assert isinstance(get_extractor_for_type("docx"), DocxExtractor)

    with pytest.raises(ValidationError):
        get_extractor_for_type("unknown_format")
