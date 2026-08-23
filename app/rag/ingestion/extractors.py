"""Multi-format text and structure extractors for PDF, Markdown, TXT, and DOCX."""

import io
import os
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Any

from pypdf import PdfReader

from app.core.exceptions import ValidationError
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ExtractedSection:
    """A structural section of a document (page, chapter, or heading block)."""

    content: str
    page_number: int | None = None
    heading: str | None = None
    section_index: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractedDocument:
    """Internal canonical representation of an extracted document."""

    title: str
    full_text: str
    sections: list[ExtractedSection]
    total_pages: int = 1
    metadata: dict[str, Any] = field(default_factory=dict)


class BaseExtractor(ABC):
    """Abstract base class for format-specific text extractors."""

    @abstractmethod
    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        """Extract structured text and section metadata from raw file bytes."""
        pass


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents preserving page boundaries."""

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            reader = PdfReader(io.BytesIO(content))
            total_pages = len(reader.pages)
            sections: list[ExtractedSection] = []
            full_text_parts: list[str] = []

            for idx, page in enumerate(reader.pages):
                page_num = idx + 1
                page_text = page.extract_text() or ""
                # Clean extra whitespace
                cleaned_page_text = re.sub(r"\s+", " ", page_text).strip()

                if cleaned_page_text:
                    full_text_parts.append(cleaned_page_text)
                    sections.append(
                        ExtractedSection(
                            content=cleaned_page_text,
                            page_number=page_num,
                            heading=f"Page {page_num}",
                            section_index=idx,
                        )
                    )

            full_text = "\n\n".join(full_text_parts)
            title = os.path.splitext(filename)[0].replace("_", " ").title()

            return ExtractedDocument(
                title=title,
                full_text=full_text,
                sections=sections,
                total_pages=total_pages,
                metadata={"total_pages": total_pages, "extractor": "pypdf"},
            )
        except Exception as e:
            logger.error("Failed to extract PDF '%s': %s", filename, str(e))
            raise ValidationError(f"Could not parse PDF document: {str(e)}") from e


class MarkdownExtractor(BaseExtractor):
    """Extractor for Markdown documents preserving heading hierarchies."""

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            text = content.decode("utf-8", errors="replace")
            lines = text.splitlines()

            sections: list[ExtractedSection] = []
            current_heading: str | None = None
            current_lines: list[str] = []
            section_idx = 0
            title = os.path.splitext(filename)[0].replace("_", " ").title()

            heading_pattern = re.compile(r"^(#{1,6})\s+(.+)$")

            for line in lines:
                match = heading_pattern.match(line.strip())
                if match:
                    # Save previous section if exists
                    if current_lines:
                        section_content = "\n".join(current_lines).strip()
                        if section_content:
                            sections.append(
                                ExtractedSection(
                                    content=section_content,
                                    heading=current_heading or "Introduction",
                                    section_index=section_idx,
                                )
                            )
                            section_idx += 1
                        current_lines = []

                    current_heading = match.group(2).strip()
                    # Set document title to first H1 if available
                    if match.group(1) == "#" and title == os.path.splitext(filename)[0].replace("_", " ").title():
                        title = current_heading
                    current_lines.append(line)
                else:
                    current_lines.append(line)

            # Flush remaining section
            if current_lines:
                section_content = "\n".join(current_lines).strip()
                if section_content:
                    sections.append(
                        ExtractedSection(
                            content=section_content,
                            heading=current_heading or "Main Content",
                            section_index=section_idx,
                        )
                    )

            if not sections and text.strip():
                sections.append(
                    ExtractedSection(
                        content=text.strip(),
                        heading="Document Content",
                        section_index=0,
                    )
                )

            return ExtractedDocument(
                title=title,
                full_text=text.strip(),
                sections=sections,
                total_pages=1,
                metadata={"format": "markdown"},
            )
        except Exception as e:
            logger.error("Failed to extract Markdown '%s': %s", filename, str(e))
            raise ValidationError(f"Could not parse Markdown document: {str(e)}") from e


class PlainTextExtractor(BaseExtractor):
    """Extractor for plain text documents."""

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            text = content.decode("utf-8", errors="replace").strip()
            title = os.path.splitext(filename)[0].replace("_", " ").title()

            # Split paragraphs as sections
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
            sections = [
                ExtractedSection(
                    content=p,
                    heading=f"Paragraph {idx + 1}",
                    section_index=idx,
                )
                for idx, p in enumerate(paragraphs)
            ]

            if not sections and text:
                sections = [ExtractedSection(content=text, heading="Content", section_index=0)]

            return ExtractedDocument(
                title=title,
                full_text=text,
                sections=sections,
                total_pages=1,
                metadata={"format": "text"},
            )
        except Exception as e:
            logger.error("Failed to extract Plain Text '%s': %s", filename, str(e))
            raise ValidationError(f"Could not parse Plain Text document: {str(e)}") from e


class DocxExtractor(BaseExtractor):
    """Extractor for Microsoft Word DOCX documents."""

    def extract(self, content: bytes, filename: str) -> ExtractedDocument:
        try:
            import docx

            doc = docx.Document(io.BytesIO(content))
            sections: list[ExtractedSection] = []
            full_text_parts: list[str] = []
            current_heading: str | None = None
            current_paras: list[str] = []
            section_idx = 0
            title = os.path.splitext(filename)[0].replace("_", " ").title()

            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue

                if p.style and p.style.name.startswith("Heading"):
                    if current_paras:
                        section_content = "\n".join(current_paras).strip()
                        if section_content:
                            sections.append(
                                ExtractedSection(
                                    content=section_content,
                                    heading=current_heading or "Overview",
                                    section_index=section_idx,
                                )
                            )
                            section_idx += 1
                        current_paras = []
                    current_heading = text
                current_paras.append(text)
                full_text_parts.append(text)

            if current_paras:
                section_content = "\n".join(current_paras).strip()
                if section_content:
                    sections.append(
                        ExtractedSection(
                            content=section_content,
                            heading=current_heading or "Content",
                            section_index=section_idx,
                        )
                    )

            full_text = "\n\n".join(full_text_parts)
            return ExtractedDocument(
                title=title,
                full_text=full_text,
                sections=sections,
                total_pages=1,
                metadata={"format": "docx"},
            )
        except Exception as e:
            logger.error("Failed to extract DOCX '%s': %s", filename, str(e))
            raise ValidationError(f"Could not parse DOCX document: {str(e)}") from e


def get_extractor_for_type(source_type: str) -> BaseExtractor:
    """Factory retrieving the appropriate extractor instance."""
    extractors = {
        "pdf": PDFExtractor(),
        "markdown": MarkdownExtractor(),
        "text": PlainTextExtractor(),
        "docx": DocxExtractor(),
    }
    if source_type not in extractors:
        raise ValidationError(f"No extractor registered for source type '{source_type}'")
    return extractors[source_type]
